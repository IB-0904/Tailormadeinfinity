from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def generate_harvard_docx(data: dict) -> bytes:
    """Generate a Harvard-style ATS-friendly resume in DOCX format."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def set_font(run, size=10, bold=False, italic=False, font_name='Arial'):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    # --- Header (Name and Contact) ---
    name = data.get("name", "Applicant Name").upper()
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(name)
    set_font(r_name, size=14, bold=True)

    contact = data.get("contact_info", "")
    if contact:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_contact = p_contact.add_run(contact)
        set_font(r_contact, size=10)

    # Helper function for section headers
    def add_section_header(title):
        p = doc.add_paragraph()
        r = p.add_run(title.upper())
        set_font(r, size=11, bold=True)
        # Add a bottom border (using an underline hack as python-docx lacks native borders)
        r.underline = True

    # --- Summary ---
    summary = data.get("summary", "")
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph()
        r = p.add_run(summary)
        set_font(r, size=10)

    # --- Experience ---
    experiences = data.get("experience", [])
    if experiences:
        add_section_header("Experience")
        for exp in experiences:
            # Company and Location (Line 1)
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            
            # Use tabs for right alignment
            p1.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))
            
            r_comp = p1.add_run(exp.get("company", ""))
            set_font(r_comp, size=10, bold=True)
            
            r_loc = p1.add_run(f'\t{exp.get("location", "")}')
            set_font(r_loc, size=10)

            # Title and Dates (Line 2)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))
            
            r_title = p2.add_run(exp.get("title", ""))
            set_font(r_title, size=10, italic=True)
            
            r_dates = p2.add_run(f'\t{exp.get("dates", "")}')
            set_font(r_dates, size=10)

            # Bullets
            for bullet in exp.get("bullets", []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                r_bullet = p_bullet.add_run(bullet)
                set_font(r_bullet, size=10)

    # --- Education ---
    education = data.get("education", [])
    if education:
        add_section_header("Education")
        for edu in education:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))
            
            r_inst = p1.add_run(edu.get("institution", ""))
            set_font(r_inst, size=10, bold=True)
            
            r_loc = p1.add_run(f'\t{edu.get("location", "")}')
            set_font(r_loc, size=10)

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.tab_stops.add_tab_stop(Inches(7.5))
            
            r_deg = p2.add_run(edu.get("degree", ""))
            set_font(r_deg, size=10, italic=True)
            
            r_dates = p2.add_run(f'\t{edu.get("dates", "")}')
            set_font(r_dates, size=10)

            for bullet in edu.get("bullets", []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                r_bullet = p_bullet.add_run(bullet)
                set_font(r_bullet, size=10)

    # --- Skills ---
    skills = data.get("skills", "")
    if skills:
        add_section_header("Skills & Additional Information")
        p = doc.add_paragraph()
        r = p.add_run(skills)
        set_font(r, size=10)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()